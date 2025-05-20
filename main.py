import tkinter as tk
from tkinter import messagebox
from abc import ABC, abstractmethod
from docx import Document
from openpyxl import Workbook

class Recipe(ABC):
    def __init__(self, ingredients):
        self.ingredients = ingredients
        self.cost = self.calculate_cost()

    @abstractmethod
    def calculate_cost(self):
        pass

    @abstractmethod
    def calculate_nutrition(self):
        pass

    def save_to_docx(self, filename):
        doc = Document()
        doc.add_heading('Recipe Report', level=1)
        doc.add_paragraph(f'Recipe: {self.__class__.__name__}')
        doc.add_paragraph(f'Ingredients: {", ".join(self.ingredients)}')
        doc.add_paragraph(f'Cost: {self.cost}')
        doc.save(filename)

    def save_to_xlsx(self, filename):
        wb = Workbook()
        ws = wb.active
        ws.title = "Recipe Report"
        ws.append(['Recipe', 'Ingredients', 'Cost'])
        ws.append([self.__class__.__name__, ', '.join(self.ingredients), self.cost])
        wb.save(filename)


class WokRecipe(Recipe):
    def calculate_cost(self):
        return len(self.ingredients) * 2

    def calculate_nutrition(self):
        return len(self.ingredients) * 100

class BurgerRecipe(Recipe):
    def calculate_cost(self):
        return len(self.ingredients) * 3

    def calculate_nutrition(self):
        return len(self.ingredients) * 150


class PizzaRecipe(Recipe):
    def calculate_cost(self):
        return len(self.ingredients) * 4

    def calculate_nutrition(self):
        return len(self.ingredients) * 200

def calculate_recipe():
    recipe_type = recipe_var.get()
    ingredients = ingredients_entry.get().split(',')
    ingredients = [ingredient.strip() for ingredient in ingredients]

    if recipe_type == 'Wok':
        recipe = WokRecipe(ingredients)
    elif recipe_type == 'Burger':
        recipe = BurgerRecipe(ingredients)
    elif recipe_type == 'Pizza':
        recipe = PizzaRecipe(ingredients)
    else:
        messagebox.showerror("Error", "Invalid recipe type")
        return

    result_text.set(f'Cost: {recipe.cost}, Nutrition: {recipe.calculate_nutrition()}')


root = tk.Tk()
root.title("Recipe Calculator")

recipe_var = tk.StringVar(value='Wok')
ingredients_entry = tk.Entry(root)
ingredients_entry.pack()

tk.Radiobutton(root, text='Wok', variable=recipe_var, value='Wok').pack()
tk.Radiobutton(root, text='Burger', variable=recipe_var, value='Burger').pack()
tk.Radiobutton(root, text='Pizza', variable=recipe_var, value='Pizza').pack()

calculate_button = tk.Button(root, text='Calculate', command=calculate_recipe)
calculate_button.pack()

result_text = tk.StringVar()
result_label = tk.Label(root, textvariable=result_text)
result_label.pack()

root.mainloop()
